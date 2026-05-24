"""
DOCX Parser
=============
Extracts text from Microsoft Word documents (.docx / .doc).
Includes:
  - Paragraph extraction (with heading detection)
  - Table extraction (formatted as readable rows)
"""

import logging
import docx
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def parse_docx(file_path: str) -> str:
    """
    Extracts text from a DOCX document including paragraphs and tables.

    Args:
        file_path: Path to the .docx file.

    Returns:
        Extracted text as a string.
    """
    result_parts = []
    try:
        doc = docx.Document(file_path)

        # Extract paragraphs with heading markers
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Prefix headings with markdown-style markers for chunking context
            style_name = para.style.name.lower() if para.style else ""
            if "heading 1" in style_name:
                result_parts.append(f"\n# {text}\n")
            elif "heading 2" in style_name:
                result_parts.append(f"\n## {text}\n")
            elif "heading 3" in style_name:
                result_parts.append(f"\n### {text}\n")
            else:
                result_parts.append(text)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            result_parts.append(f"\n[Table {table_idx + 1}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    result_parts.append(" | ".join(cells))
            result_parts.append("")

    except Exception as e:
        logger.error(f"Error parsing DOCX '{file_path}': {e}")

    return "\n".join(result_parts)
